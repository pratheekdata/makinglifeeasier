import os
from docx import Document
from docx.shared import Pt

def add_directory_to_doc(doc, path, level=0):
    """
    Recursively add directory contents to the Word document.
    """
    indent = ' ' * (level * 4)
    for item in os.listdir(path):
        item_path = os.path.join(path, item)
        if os.path.isdir(item_path):
            doc.add_paragraph(f"{indent}{item}/", style='Heading2')
            add_directory_to_doc(doc, item_path, level + 1)
        else:
            doc.add_paragraph(f"{indent}{item}")

def create_directory_hierarchy_doc(root_dir, output_file):
    """
    Create a Word document with the directory hierarchy of the given root directory.
    """
    doc = Document()
    doc.add_heading(f'Directory Hierarchy of {root_dir}', 0)

    add_directory_to_doc(doc, root_dir)

    doc.save(output_file)
    print(f"Document saved as {output_file}")

if __name__ == "__main__":
    root_directory = input("Enter the path of the directory: ")
    output_filename = input("Enter the output Word document filename (with .docx extension): ")
    create_directory_hierarchy_doc(root_directory, output_filename)
